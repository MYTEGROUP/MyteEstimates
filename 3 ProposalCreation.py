#proposalCreation.py
import json
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement


def parse_json_to_table(file_path, feature_list_id):
    with open(file_path, 'r') as file:
        all_data = json.load(file)

    # Filter for the specified FeatureListID
    data = next((item for item in all_data if item["FeatureListID"] == feature_list_id), None)
    if not data:
        raise ValueError(f"No data found for FeatureListID: {feature_list_id}")

    rows = []
    total_hours_per_feature = {}
    for item in data["Features"]:
        feature = item["Feature"]
        sub_feature = item["SubFeature"]
        task_name = item["TaskName"]
        task_description = item["TaskDescription"]["Description"]
        estimated_hours = item["TaskDescription"]["EstimatedHours"]
        if feature not in total_hours_per_feature:
            total_hours_per_feature[feature] = 0
        total_hours_per_feature[feature] += estimated_hours
        rows.append({
            "Feature Name": feature,
            "Sub Feature": sub_feature,
            "Task Name": task_name,
            "Task Description": task_description,
            "Estimated Hours": estimated_hours
        })

    df = pd.DataFrame(rows)
    final_df = pd.DataFrame()
    for feature, tasks in df.groupby("Feature Name"):
        tasks_without_feature_name = tasks.copy()
        tasks_without_feature_name["Feature Name"] = ""
        final_df = pd.concat([final_df, tasks.iloc[0:1]], ignore_index=True)
        final_df = pd.concat([final_df, tasks_without_feature_name.iloc[1:]], ignore_index=True)
        total_hours_row = pd.DataFrame([{
            "Feature Name": feature,
            "Sub Feature": "",
            "Task Name": "",
            "Task Description": "Total Hours",
            "Estimated Hours": total_hours_per_feature[feature]
        }])
        final_df = pd.concat([final_df, total_hours_row], ignore_index=True)
    return final_df

def set_cell_border(cell, border_color_hex="000000", border_width=Pt(1)):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(int(border_width.pt)))
        border.set(qn('w:color'), border_color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_column_widths(table, widths):
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)

def set_margins(doc, left, right):
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def calculate_column_widths(table):
    page_width = 8.5
    left_margin = 0.5
    right_margin = 0.5
    usable_width = page_width - left_margin - right_margin
    num_columns = len(table.columns)
    column_width = usable_width / num_columns
    return [column_width for _ in range(num_columns)]

def set_header(doc, logo_path):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]

    # Add logo to the header
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(2))  # Adjust the width as needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def save_to_word(df, file_path):
    doc = Document()
    set_margins(doc, 0.5, 0.5)

    # Set header with logo
    logo_path = 'static/Logo.png'
    set_header(doc, logo_path)

    # Add a blank first page with space after the header
    doc.add_paragraph("\n")  # Adds a line break for spacing

    # Proposal text or other content for the first page goes here...

    doc.add_page_break()

    # Add Annexe A and Hours Breakdown text before the table with space after the header
    doc.add_paragraph("\n")  # Adds a line break for spacing
    annexe_a = doc.add_paragraph("Annexe A")
    annexe_a.style = doc.styles['Heading 1']
    annexe_a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    hours_breakdown = doc.add_paragraph("Hours Breakdown")
    hours_breakdown.style = doc.styles['Heading 2']
    hours_breakdown.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create the table
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    widths = calculate_column_widths(table)
    set_column_widths(table, widths)

    # Apply header row formatting
    for j in range(df.shape[-1]):
        cell = table.cell(0, j)
        cell.text = df.columns[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, "000000", Pt(1))
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w'))))
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Apply cell formatting and set row properties to prevent splitting
    for i in range(df.shape[0]):
        row = table.rows[i + 1]
        row.allow_break_across_pages = False  # Prevent row from splitting across pages
        for j in range(df.shape[-1]):
            cell = row.cells[j]
            cell.text = str(df.iloc[i, j])
            set_cell_border(cell, "000000", Pt(1))
            # Center-align 'Estimated Hours' column
            if j == 4:  # Assuming 'Estimated Hours' is the 5th column (index 4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Special formatting for 'Total Hours' row
            if df.iloc[i, 3] == "Total Hours" and df.iloc[i, 0] != "":
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w'))))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

    # Add an empty paragraph at the end of the table for spacing
    doc.add_paragraph("\n")

    doc.save(file_path)


# Use the function
json_file_path = 'storage/HourEstimate.json'
word_file_path = 'Proposals/Proposal.docx'

# Accept FeatureListID as console input
try:
    feature_list_id_to_process = int(input("Enter the FeatureListID: "))
except ValueError:
    print("Invalid input. Please enter a numeric FeatureListID.")
    exit(1)

df = parse_json_to_table(json_file_path, feature_list_id_to_process)
save_to_word(df, word_file_path)